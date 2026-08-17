"""简历文本提取：PDF / DOCX / TXT。

所有解析均从内存字节进行，文件不落盘，符合「个人信息仅保留在本地」的隐私要求。
"""
from __future__ import annotations

from io import BytesIO


class ParseError(Exception):
    pass


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ParseError("缺少 pypdf 依赖，请运行 pip install pypdf")
    try:
        reader = PdfReader(BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise ParseError(f"PDF 解析失败：{e}")
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    text = "\n".join(pages).strip()
    if not text:
        raise ParseError("未能从 PDF 中提取到文本（可能是扫描件/图片型 PDF）。请改用「直接粘贴简历文本」。")
    return text


def _parse_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise ParseError("缺少 python-docx 依赖，请运行 pip install python-docx")
    try:
        document = docx.Document(BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise ParseError(f"Word 解析失败：{e}")
    parts = []
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ParseError("Word 文档中未提取到文本内容。")
    return text


def _parse_txt(data: bytes) -> str:
    for enc in ("utf-8", "gb18030", "utf-16"):
        try:
            return data.decode(enc).strip()
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="ignore").strip()


_SUPPORTED = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".doc": None,  # 老版 .doc 二进制格式，需转换，见注释
    ".txt": _parse_txt,
    ".md": _parse_txt,
}


def extract_text(filename: str, data: bytes) -> str:
    """按扩展名提取文本；不支持的格式抛出 ParseError。"""
    ext = ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""
    parser = _SUPPORTED.get(ext)
    if parser is None:
        if ext == ".doc":
            raise ParseError("暂不支持老版 .doc 格式，请在 Word 中另存为 .docx 或 .txt 后再上传。")
        raise ParseError(f"不支持的文件格式「{ext or filename}」，支持 PDF / DOCX / TXT。")
    text = parser(data)
    if not text.strip():
        raise ParseError("文件内容为空。")
    return text
